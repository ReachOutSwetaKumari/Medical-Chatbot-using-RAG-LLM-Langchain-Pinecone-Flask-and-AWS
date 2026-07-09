from langchain_community.document_loaders import PyMuPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document
import os


def load_pdf_data(data_path):
    files = [f for f in os.listdir(data_path) if f.endswith('.pdf')]
    if not files:
        print("No PDF files found inside the target folder path!")
        return []
    target_file = os.path.join(data_path, files[0])
    print(f"Targeting file directly: {target_file}")
    loader = PyMuPDFLoader(target_file)
    documents = loader.load()
    return documents


def load_single_pdf(file_path):
    loader = PyMuPDFLoader(file_path)
    return loader.load()


def load_docx_file(file_path):
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text.append(row_text)
        text = "\n".join(full_text)
        if not text.strip():
            raise Exception("No readable text found in DOCX file.")
        return [Document(page_content=text, metadata={"source": file_path, "type": "docx"})]
    except ImportError:
        raise Exception("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        raise Exception(f"DOCX loading failed: {str(e)}")


def load_csv_file(file_path):
    try:
        import pandas as pd
        df = pd.read_csv(file_path, encoding='utf-8', on_bad_lines='skip')
        text = (
            f"CSV Data — {len(df)} rows, {len(df.columns)} columns.\n"
            f"Columns: {', '.join(str(c) for c in df.columns.tolist())}\n\n"
            f"{df.to_string(index=False)}"
        )
        return [Document(page_content=text, metadata={"source": file_path, "type": "csv"})]
    except ImportError:
        raise Exception("pandas not installed. Run: pip install pandas")
    except Exception as e:
        raise Exception(f"CSV loading failed: {str(e)}")


def load_xlsx_file(file_path):
    try:
        import pandas as pd
        xl = pd.ExcelFile(file_path, engine='openpyxl')
        all_docs = []
        for sheet_name in xl.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name, engine='openpyxl')
            text = f"Sheet: {sheet_name}\nRows: {len(df)}, Columns: {len(df.columns)}\n\n{df.to_string(index=False)}"
            all_docs.append(Document(
                page_content=text,
                metadata={"source": file_path, "type": "xlsx", "sheet": str(sheet_name)}
            ))
        return all_docs if all_docs else [Document(page_content="Empty Excel file", metadata={"source": file_path, "type": "xlsx"})]
    except ImportError:
        raise Exception("pandas/openpyxl not installed. Run: pip install pandas openpyxl")
    except Exception as e:
        raise Exception(f"XLSX loading failed: {str(e)}")


def load_image_with_ocr(file_path):
    try:
        import easyocr
        reader = easyocr.Reader(['en'], verbose=False)
        results = reader.readtext(file_path, detail=0)
        text = "\n".join(str(r) for r in results).strip()
        if not text:
            text = "No readable text detected in image."
        return [Document(page_content=text, metadata={"source": file_path, "type": "image_ocr"})]
    except ImportError:
        raise Exception("easyocr not installed. Run: pip install easyocr")
    except Exception as e:
        raise Exception(f"OCR processing failed: {str(e)}")


def load_file_by_type(file_path):
    ext = os.path.splitext(file_path)[1].lower().lstrip('.')
    if ext == 'pdf':
        return load_single_pdf(file_path)
    elif ext == 'docx':
        return load_docx_file(file_path)
    elif ext == 'csv':
        return load_csv_file(file_path)
    elif ext in ('xlsx', 'xls'):
        return load_xlsx_file(file_path)
    elif ext in ('png', 'jpg', 'jpeg', 'tiff', 'bmp', 'gif', 'webp'):
        return load_image_with_ocr(file_path)
    else:
        raise Exception(f"Unsupported file type: .{ext}")


def text_split(extracted_data):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1300, chunk_overlap=200)
    text_chunks = text_splitter.split_documents(extracted_data)
    return text_chunks


def download_hugging_face_embeddings():
    embeddings = HuggingFaceEmbeddings(
        model_name="BAAI/bge-base-en-v1.5",
        encode_kwargs={"normalize_embeddings": True, "batch_size": 64}
    )
    return embeddings
